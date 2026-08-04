"""
Index a folder of documents so the assistant can answer from them.

    uv run python -m app.ingest ~/notes
    uv run python -m app.ingest ~/notes --ext .md .txt

Re-running is safe: each file's chunks are keyed by its path, and
ingest_document deletes a source's old chunks before writing new ones, so
an edited file is replaced rather than duplicated.

Text files only for now. PDF and docx need an extractor -- pypdf and
python-docx are the usual choices, and both are a bigger job than they
look because layout, tables and columns all affect how the text chunks.
"""

import argparse
import asyncio
import io
import pathlib

import docx
import pypdf

from app.db import DEFAULT_TENANT_ID
from app.rag import ingest_document

TEXT_EXTENSIONS = (".md", ".txt", ".markdown", ".rst")
DEFAULT_EXTENSIONS = TEXT_EXTENSIONS + (".pdf", ".docx")


def extract_text(data: bytes, filename: str) -> str:
    """Text from raw bytes, dispatched on the filename's extension.

    Bytes rather than a path because the same three formats arrive from
    three places now: the filesystem, a WhatsApp media download, and a
    Drive export. Only the disk case has a path.
    """
    suffix = pathlib.Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(io.BytesIO(data))
    if suffix == ".docx":
        return _read_docx(io.BytesIO(data))
    return data.decode("utf-8", errors="replace")


def _read_pdf(path) -> str:
    """Page text, joined with blank lines so chunking can split on them.

    Extraction is best-effort by nature: a PDF stores glyphs at
    coordinates, not sentences. Multi-column layouts interleave, tables
    lose their structure, and scanned pages yield nothing at all without
    OCR. A page that comes back empty is skipped rather than embedded --
    an empty chunk matches everything weakly and pollutes results.
    """
    pages = []
    reader = pypdf.PdfReader(path)
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _read_docx(path) -> str:
    """Paragraphs and table cells.

    python-docx keeps tables out of document.paragraphs, so a file whose
    content is mostly tabular extracts as almost nothing unless they are
    walked separately.
    """
    document = docx.Document(path)
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def read_document(path: pathlib.Path) -> str:
    """Disk convenience wrapper -- one extraction path, not two."""
    return extract_text(path.read_bytes(), path.name)


async def ingest_folder(folder: pathlib.Path, extensions: tuple[str, ...]) -> None:
    files = sorted(
        p for p in folder.rglob("*") if p.is_file() and p.suffix.lower() in extensions
    )
    if not files:
        print(f"no {'/'.join(extensions)} files under {folder}")
        return

    total = 0
    for path in files:
        try:
            content = read_document(path)
        except Exception as exc:
            # One unreadable file shouldn't abandon the whole folder --
            # encrypted PDFs and malformed docx are common in any real
            # directory.
            print(f"  skipped {path.name}: {type(exc).__name__}")
            continue

        if not content.strip():
            # Usually a scanned PDF: real pages, no extractable text.
            # Needs OCR, which is a different problem.
            print(f"  skipped {path.name}: no extractable text")
            continue

        chunks = await ingest_document(
            DEFAULT_TENANT_ID,
            # Full path as the key: two files can share a name in
            # different folders, and re-ingesting must replace the right one.
            source_id=str(path.resolve()),
            source_name=path.stem,
            content=content,
        )
        total += chunks
        print(f"  {path.name:<40} {chunks:>3} chunks")

    print(f"\nindexed {len(files)} files, {total} chunks")


def main() -> None:
    parser = argparse.ArgumentParser(description="Index documents for RAG.")
    parser.add_argument("folder", type=pathlib.Path)
    parser.add_argument("--ext", nargs="+", default=list(DEFAULT_EXTENSIONS))
    args = parser.parse_args()

    if not args.folder.is_dir():
        raise SystemExit(f"not a directory: {args.folder}")

    asyncio.run(
        ingest_folder(args.folder, tuple(e.lower() for e in args.ext))
    )


if __name__ == "__main__":
    main()
